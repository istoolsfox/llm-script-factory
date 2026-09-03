"""
Stage 3 Service: Scene Writer (Episode Outlines)

Converts Stage 2 outlines into detailed episode outlines with scene breakdowns.
"""
import os
import json
from typing import List, Dict, Optional, Any
from services.base import BaseService


class Stage3Service(BaseService):
    """
    Service Controller for Stage 3: Scene Writer.
    Converts Stage 2 outlines into detailed episode outlines with scene breakdowns.
    """
    
    def __init__(self):
        super().__init__()
        
        # === Prompt Templates ===
        self.TMPL_SYS = "stage3/1_context_sys.j2"
        self.TMPL_DTG = "stage3/2_context_dtg.j2"
        self.TMPL_USER = "stage3/3_episode_user.j2"
        
        # Schema (for validation)
        self.SCHEMA_EPISODE = "prompts/stage3/schema_episode.json"
        
        # DTG Files (Pruned: Core + Dialog + Sample)
        self.DTG_FILES = [
            "dtgCore_2025_1222_0001.md",
            "dtgDialog_2025_1223_0001.md",
            "dtgSample_2025_1229_0001.md",
        ]
        
        # API Response Schema (for structured output enforcement)
        self.RESPONSE_SCHEMA = {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "ep_id": {"type": "integer"},
                    "title": {"type": "string"},
                    "scenes": {
                        "type": "array",
                        "items": {
                            "type": "object",
                            "properties": {
                                "scene_id": {"type": "string"},
                                "location": {"type": "string"},
                                "content": {"type": "string"}
                            },
                            "required": ["scene_id", "location", "content"]
                        }
                    },
                    "hook": {"type": "string"}
                },
                "required": ["ep_id", "title", "scenes", "hook"]
            }
        }

    # =========================================================================
    # DTG CONTENT
    # =========================================================================
    def _prepare_dtg_content(self) -> str:
        """Load DTG Core + Dialog."""
        return self.prompts.load_dtg_theory(branch="dtg/Distill-1", file_list=self.DTG_FILES)

    # =========================================================================
    # GENERATION
    # =========================================================================
    def generate_batch(
        self,
        project_name: str,
        start_ep: int,
        end_ep: int,
        temperature: float = 0.7
    ) -> List[Dict]:
        """
        Generate detailed episode outlines for a batch.
        """
        # 1. Load required data
        story_bible = self._load_story_bible(project_name)
        s2_outlines = self._load_s2_outlines(project_name)
        previous_s3_outlines = self._load_s3_outlines(project_name)

        if not s2_outlines:
            raise ValueError("请先完成 Stage 2（结构构建），生成分集大纲后再来编写集纲")

        # 2. Prepare user prompt (always needed)
        context_window = self._get_context_window(start_ep, s2_outlines)
        rearview_mirror = self._get_rearview_mirror(previous_s3_outlines, start_ep)

        context_str = json.dumps(context_window, indent=2, ensure_ascii=False)
        rearview_str = json.dumps(rearview_mirror, indent=2, ensure_ascii=False)

        user_content = self.prompts.render(
            self.TMPL_USER,
            context_window=context_str,
            rearview_mirror=rearview_str,
            start_ep=start_ep,
            end_ep=end_ep
        )

        # 3. Resolve model from project settings
        model_key, temperature = self.resolve_model_key(project_name, "stage3")

        # 4. Generate
        return self._generate_batch_raw(model_key, user_content, story_bible, project_name, temperature)

    def _generate_batch_raw(
        self, 
        model_key: str, 
        user_content: str,
        story_bible: Dict,
        project_name: str,
        temperature: float
    ) -> List[Dict]:
        """Raw path: full system + context + user."""
        story_bible_str = json.dumps(story_bible, indent=2, ensure_ascii=False)
        user_input = self._load_user_input(project_name)
        dtg_raw = self._prepare_dtg_content()
        
        sys_content = self.prompts.render(self.TMPL_SYS, full_context="", story_bible="")
        dtg_content = self.prompts.render(
            self.TMPL_DTG,
            full_context=dtg_raw,
            story_bible=story_bible_str,
            user_input=user_input
        )
        
        result = self.process_request(
            model_key=model_key,
            user_content=user_content,
            system_content=sys_content,
            context_content=dtg_content,
            temperature=temperature,
            source="stage3/generate/raw",
            response_schema=self.RESPONSE_SCHEMA
        )
        return self._normalize_result(result)

    # =========================================================================
    # DATA LOADING
    # =========================================================================
    def _load_story_bible(self, project_name: str) -> Dict:
        """Load story bible from Stage 1."""
        project_dir = self.projects.get_project_path(project_name)
        path = os.path.join(project_dir, "1_ideas/story_bible.json")
        return self.files.load_json(path, default={})

    def _load_user_input(self, project_name: str) -> str:
        """Load user input (concept) from Stage 1."""
        project_dir = self.projects.get_project_path(project_name)
        if not project_dir:
            return ""
        path = os.path.join(project_dir, "1_ideas/user_input.json")
        data = self.files.load_json(path, default={})
        return data.get("concept", "")

    def _load_s2_outlines(self, project_name: str) -> List[Dict]:
        """Load Stage 2 outlines."""
        project_dir = self.projects.get_project_path(project_name)
        path = os.path.join(project_dir, "2_structure/detailed_outlines.json")
        return self.files.load_json(path, default=[])

    def _load_s3_outlines(self, project_name: str) -> List[Dict]:
        """Load existing Stage 3 outlines."""
        project_dir = self.projects.get_project_path(project_name)
        path = os.path.join(project_dir, "3_scripts/episode_outlines.json")
        return self.files.load_json(path, default=[])

    # =========================================================================
    # HELPERS
    # =========================================================================
    def _get_context_window(self, anchor_ep: int, s2_outlines: List[Dict]) -> List[Dict]:
        """Get Anchor - 5 to Anchor + 5 from Stage 2 outlines."""
        start_range = max(1, anchor_ep - 5)
        end_range = min(80, anchor_ep + 5)
        return [ep for ep in s2_outlines if start_range <= ep.get('ep_id', 0) <= end_range]

    def _get_rearview_mirror(self, s3_outlines: List[Dict], start_ep: int) -> List[Dict]:
        """Get last 3 generated detailed outlines BEFORE the current batch."""
        # Filter to only episodes before the current batch
        previous_eps = [ep for ep in s3_outlines if ep.get('ep_id', 0) < start_ep]
        sorted_list = sorted(previous_eps, key=lambda x: x.get('ep_id', 0))
        return sorted_list[-3:] if sorted_list else []

    def _normalize_result(self, result: Any) -> List[Dict]:
        """Normalize LLM result to ensure it's a list of episodes."""
        if isinstance(result, dict):
            if "episodes" in result:
                return result["episodes"]
            for v in result.values():
                if isinstance(v, list):
                    return v
            return [result]
        elif isinstance(result, list):
            return result
        return []

    # =========================================================================
    # SAVING
    # =========================================================================
    def save_batch(self, project_name: str, new_batch: List[Dict]) -> bool:
        """Merge and save to 3_scripts/episode_outlines.json."""
        project_dir = self.projects.get_project_path(project_name)
        path = os.path.join(project_dir, "3_scripts/episode_outlines.json")
        
        # 1. Load existing
        existing_data = self.files.load_json(path, default=[])
        
        # 2. Merge (Upsert by ep_id)
        data_map = {item['ep_id']: item for item in existing_data if 'ep_id' in item}
        for item in new_batch:
            ep_id = item.get('ep_id')
            if ep_id:
                data_map[ep_id] = item
                
        merged = list(data_map.values())
        merged.sort(key=lambda x: x.get('ep_id', 0))
        
        # 3. Save
        return self.files.save_json(path, merged)

    def clear_all_scripts(self, project_name: str) -> bool:
        """
        Clear all episode outlines for a project.
        Writes an empty array to 3_scripts/episode_outlines.json.
        """
        project_dir = self.projects.get_project_path(project_name)
        path = os.path.join(project_dir, "3_scripts/episode_outlines.json")
        return self.files.save_json(path, [])

    # =========================================================================
    # EXPORT DOCX
    # =========================================================================
    def export_docx(self, project_name: str) -> str:
        """
        Export episode outlines to DOCX file with StoryBible header.
        
        Format:
        剧名：XXXXX
        简介：XXXXX
        人设表：...
        前3卡卡纲：...
        (换页)
        第N集...
        
        Returns: Path to generated DOCX file.
        """
        from docx import Document
        from docx.shared import Pt
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        
        project_dir = self.projects.get_project_path(project_name)
        if not project_dir:
            raise ValueError(f"Project '{project_name}' not found")
        
        # Load data
        outlines = self._load_s3_outlines(project_name)
        if not outlines:
            raise ValueError("No episode outlines found to export")
        
        story_bible = self._load_story_bible(project_name)
        synopsis = story_bible.get('synopsis', {})
        rough_skeleton = story_bible.get('rough_skeleton', [])
        
        # Sort outlines by ep_id
        outlines.sort(key=lambda x: x.get('ep_id', 0))
        
        # Create document
        doc = Document()
        
        # Set default font
        style = doc.styles['Normal']
        style.font.name = 'SimSun'
        style.font.size = Pt(12)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        
        # === Header Section (from StoryBible) ===
        
        # 剧名
        title = synopsis.get('title', project_name)
        doc.add_paragraph(f"剧名：{title}")
        doc.add_paragraph("")
        
        # 简介（使用故事主线而非一句话介绍）
        story_synopsis = synopsis.get('synopsis', synopsis.get('logline', ''))
        doc.add_paragraph(f"简介：{story_synopsis}")
        doc.add_paragraph("")
        
        # 人设表
        doc.add_paragraph("人设表：")
        characters = synopsis.get('characters', [])
        for char in characters:
            name = char.get('name', '')
            role = char.get('role', '')
            desc = char.get('desc', '')
            doc.add_paragraph(f"{name}（{role}）：{desc}")
        doc.add_paragraph("")
        
        # 前3卡卡纲
        doc.add_paragraph("前3卡卡纲：")
        for card in rough_skeleton[:3]:
            card_id = card.get('card_id', 0)
            summary = card.get('one_sentence_summary', '')
            doc.add_paragraph(f"Card {card_id}：{summary}")
        doc.add_paragraph("")
        
        # === Page Break ===
        doc.add_page_break()
        
        # === Episode Outlines ===
        for ep in outlines:
            ep_id = ep.get('ep_id', 0)
            scenes = ep.get('scenes', [])
            hook = ep.get('hook', '')
            
            # Episode header
            doc.add_paragraph(f"第{ep_id}集")
            
            # Scenes
            for scene in scenes:
                scene_id = scene.get('scene_id', '')
                location = scene.get('location', '')
                content = scene.get('content', '')
                
                doc.add_paragraph(scene_id)
                doc.add_paragraph(location)
                doc.add_paragraph(content)
            
            # Hook
            if hook:
                doc.add_paragraph(hook)
            
            # Separator
            doc.add_paragraph("")
        
        # Save
        output_dir = os.path.join(project_dir, "3_scripts")
        os.makedirs(output_dir, exist_ok=True)
        output_path = os.path.join(output_dir, f"集纲_{project_name}.docx")
        doc.save(output_path)
        
        print(f"✅ Stage3 DOCX exported: {output_path}")
        return output_path

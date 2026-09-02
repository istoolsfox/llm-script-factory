"""
Stage 5 Service: Script Polisher

Polishes Stage 4 drafts into refined production-ready scripts.
"""
import os
import json
import re
from typing import List, Dict, Optional, Any, Union, Tuple
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from services.base import BaseService


class Stage5Service(BaseService):
    """
    Service Controller for Stage 5: Script Polisher.
    Polishes Stage 4 drafts into refined production-ready scripts.
    """
    
    def __init__(self):
        super().__init__()
        
        # === Prompt Templates ===
        self.TMPL_SYS = "stage5/1_format_sys.j2"
        self.TMPL_EXAMPLES = "stage5/2_examples_dtg.j2"
        self.TMPL_USER = "stage5/3_refine_user.j2"
        
        # Schema
        self.SCHEMA_REFINED_SCRIPT = "prompts/stage5/schema_refined_script.json"

    # =========================================================================
    # GENERATION
    # =========================================================================
    def generate_batch(
        self,
        project_name: str,
        start_ep: int,
        end_ep: int,
        temperature: float = 0.3
    ) -> List[Dict]:
        """
        Generate polished scripts for a batch.
        """
        # 1. Load required data
        story_bible = self._load_story_bible(project_name)
        story_bible_str = self._get_bible_text(story_bible)
        s4_scripts = self._load_s4_scripts(project_name)

        # Filter to requested range
        batch_scripts = [ep for ep in s4_scripts if start_ep <= ep.get('ep_id', 0) <= end_ep]
        if not batch_scripts:
            raise ValueError(f"No Stage 4 scripts found for episodes {start_ep}-{end_ep}")

        # 2. Prepare user prompt (always needed)
        registry = self._load_registry(project_name)
        status_list = self._format_character_status(registry, story_bible_str)
        raw_text_block = self._format_raw_scripts(batch_scripts)

        user_content = self.prompts.render(
            self.TMPL_USER,
            raw_script=raw_text_block,
            character_status_list=status_list
        )

        # 3. Resolve model from project settings
        model_key, temperature = self.resolve_model_key(project_name, "stage5")

        # 4. Generate
        result = self._generate_batch_raw(model_key, user_content, story_bible_str, project_name, temperature)

        # 5. Post-process (update registry)
        self._update_registry_from_response(project_name, result)

        return self._normalize_result(result, batch_scripts)

    def _generate_batch_raw(
        self, 
        model_key: str, 
        user_content: str,
        story_bible_str: str,
        project_name: str,
        temperature: float
    ) -> Any:
        """Raw path: full system + context + user."""
        user_input = self._load_user_input(project_name)
        sys_content = self.prompts.render(self.TMPL_SYS)
        examples_content = self.prompts.render(self.TMPL_EXAMPLES, story_bible=story_bible_str, user_input=user_input)
        
        return self.process_request(
            model_key=model_key,
            user_content=user_content,
            system_content=sys_content,
            context_content=examples_content,
            temperature=temperature,
            source="stage5/generate/raw"
        )

    # =========================================================================
    # DATA LOADING
    # =========================================================================
    def _load_story_bible(self, project_name: str) -> Dict:
        """Load story bible from Stage 1."""
        project_dir = self.projects.get_project_path(project_name)
        path = os.path.join(project_dir, "1_ideas/story_bible.json")
        return self.files.load_json(path, default={})

    def _get_bible_text(self, story_bible: Dict) -> str:
        """Extract raw_content from story bible."""
        if isinstance(story_bible, dict):
            return story_bible.get("raw_content", json.dumps(story_bible, ensure_ascii=False))
        return ""

    def _load_user_input(self, project_name: str) -> str:
        """Load user input (concept) from Stage 1."""
        project_dir = self.projects.get_project_path(project_name)
        if not project_dir:
            return ""
        path = os.path.join(project_dir, "1_ideas/user_input.json")
        data = self.files.load_json(path, default={})
        return data.get("concept", "")

    def _load_s4_scripts(self, project_name: str) -> List[Dict]:
        """Load Stage 4 scripts."""
        project_dir = self.projects.get_project_path(project_name)
        path = os.path.join(project_dir, "4_scripts/script_drafts.json")
        return self.files.load_json(path, default=[])

    def _load_registry(self, project_name: str) -> Dict:
        """Load character appearance registry."""
        project_dir = self.projects.get_project_path(project_name)
        path = os.path.join(project_dir, "5_scripts/character_registry.json")
        return self.files.load_json(path, default={})

    def _save_registry(self, project_name: str, registry: Dict) -> bool:
        """Save character appearance registry."""
        project_dir = self.projects.get_project_path(project_name)
        path = os.path.join(project_dir, "5_scripts/character_registry.json")
        return self.files.save_json(path, registry)

    # =========================================================================
    # HELPERS
    # =========================================================================
    def _format_character_status(self, registry: Dict, story_bible: str) -> str:
        """Generate the status list for the prompt."""
        all_chars = self._extract_chars_from_bible(story_bible)
        lines = []
        for name, first_ep in registry.items():
            lines.append(f"{name}：第{first_ep}集出场，无需介绍字幕")
        for name in all_chars:
            if name not in registry:
                lines.append(f"{name}：未出场，请只在他出现的第一集增加介绍字幕。")
        if not lines:
            return "（暂无角色记录，请对所有首次出场的重要角色添加字幕）"
        return "\n".join(lines)

    def _extract_chars_from_bible(self, story_bible: str) -> List[str]:
        """Extract character names from Story Bible text."""
        if not story_bible: return []
        names = set()
        matches = re.findall(r'(?:人物|角色)[:：]\s*([^\n]+)', story_bible)
        for m in matches:
            clean = re.split(r'[（(]', m)[0].strip()
            if clean: names.add(clean)
        return list(names)

    def _format_raw_scripts(self, raw_scripts: List[Dict]) -> str:
        """Serialize Stage 4 scripts for Prompt."""
        text_block = ""
        for ep in raw_scripts:
            eid = ep.get('ep_id', '?')
            text_block += f"\n--- 第 {eid} 集 ---\n"
            if 'content' in ep and ep['content']:
                text_block += f"{ep['content']}\n"
            elif 'scenes' in ep:
                for sc in ep['scenes']:
                    sid = sc.get('scene_id', '')
                    h = sc.get('header', {})
                    time = h.get('time') or sc.get('time') or ""
                    if time: text_block += f"{sid} {time}\n"
                    else: text_block += f"{sid}\n"
                    loc = h.get('location') or sc.get('location') or ""
                    if loc: text_block += f"场景：{loc}\n"
                    chars = h.get('characters') or sc.get('characters') or ""
                    if chars: text_block += f"人物：{chars}\n"
                    body = sc.get('raw_body') or sc.get('content') or ""
                    text_block += f"\n{body}\n\n"
            elif 'text' in ep:
                text_block += f"{ep['text']}\n"
        return text_block

    def _update_registry_from_response(self, project_name: str, response_data: Any):
        """Update the registry with new appearances from this batch."""
        if isinstance(response_data, dict):
            new_apps = response_data.get("new_appearances", [])
            if not new_apps: return
            
            registry = self._load_registry(project_name)
            updated = False
            for app in new_apps:
                name = app.get("name")
                ep_id = app.get("ep_id")
                if name and name not in registry:
                    registry[name] = ep_id
                    updated = True
            if updated:
                self._save_registry(project_name, registry)

    def _normalize_result(self, result: Any, original_batch: List[Dict]) -> List[Dict]:
        """Normalize LLM result to ensure it's a list of episodes."""
        if isinstance(result, list):
            return result
        if isinstance(result, dict):
            if "episodes" in result:
                return result["episodes"]
            for v in result.values():
                if isinstance(v, list):
                    return v
            return [result]
        return original_batch

    # =========================================================================
    # SAVING
    # =========================================================================
    def save_batch(self, project_name: str, new_batch: List[Dict]) -> bool:
        """Merge and save to 5_scripts/refined_scripts.json."""
        project_dir = self.projects.get_project_path(project_name)
        path = os.path.join(project_dir, "5_scripts/refined_scripts.json")
        
        # 1. Load existing
        existing_data = self.files.load_json(path, default=[])
        
        # 2. Merge (Upsert by ep_id)
        data_map = {item['ep_id']: item for item in existing_data if 'ep_id' in item}
        for item in new_batch:
            if 'ep_id' in item:
                data_map[item['ep_id']] = item
                
        merged = list(data_map.values())
        merged.sort(key=lambda x: x.get('ep_id', 0))
        
        # 3. Save
        return self.files.save_json(path, merged)

    def clear_all_scripts(self, project_name: str) -> bool:
        """
        Clear all refined scripts for a project.
        Writes an empty array to 5_scripts/refined_scripts.json.
        Also clears character_registry.json.
        """
        project_dir = self.projects.get_project_path(project_name)
        scripts_path = os.path.join(project_dir, "5_scripts/refined_scripts.json")
        registry_path = os.path.join(project_dir, "5_scripts/character_registry.json")
        
        # Clear both files
        self.files.save_json(scripts_path, [])
        self.files.save_json(registry_path, {})
        return True

    # =========================================================================
    # IMPORT/EXPORT HELPERS (preserved from original)
    # =========================================================================
    def export_docx(self, project_name: str, scripts_data: List[Dict], output_subfolder: str = "5_scripts") -> str:
        """Export scripts to DOCX format.
        
        Args:
            project_name: Name of the project
            scripts_data: List of episode scripts to export
            output_subfolder: Subdirectory to save the file (default: 5_scripts)
        """
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'SimSun' 
        style.font.size = Pt(12)
        style.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        scripts_data.sort(key=lambda x: x.get('ep_id', 0))
        
        for idx, ep in enumerate(scripts_data):
            eid = ep.get('ep_id', '?')
            if idx > 0:
                doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            head_p = doc.add_paragraph(f"第 {eid} 集")
            head_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in head_p.runs:
                run.bold = True
                run.font.name = 'SimSun'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
                run.font.size = Pt(12)
            doc.add_paragraph()
            
            scenes = ep.get('scenes', [])
            for scene in scenes:
                def add_p(text, bold=False):
                    p = doc.add_paragraph(str(text))
                    for run in p.runs:
                        run.font.name = 'SimSun'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
                        run.font.size = Pt(12)
                        if bold: run.bold = True
                    return p
                    
                sid = scene.get('scene_id', '')
                add_p(f"{sid}", bold=True)
                time = scene.get('time', '')
                if time: add_p(time)
                loc = scene.get('location', '')
                if loc: add_p(f"场景：{loc}")
                chars = scene.get('characters', '')
                if chars: add_p(f"人物：{chars}")
                doc.add_paragraph()
                content = scene.get('content', '')
                if content: add_p(content)
                doc.add_paragraph()
                doc.add_paragraph()
                
        project_dir = self.projects.get_project_path(project_name)
        filename = f"{project_name}_script_export.docx"
        out_path = os.path.join(project_dir, output_subfolder, filename)
        os.makedirs(os.path.dirname(out_path), exist_ok=True)
        doc.save(out_path)
        return out_path

    # =========================================================================
    # COPY FROM PREVIOUS STAGE
    # =========================================================================
    def copy_from_s4(self, project_name: str) -> Dict[str, Any]:
        """
        Copy all Stage 4 files to Stage 5 (reset/initialize).
        
        Steps:
        1. Clear all files in 5_scripts folder
        2. Copy all files from 4_scripts folder to 5_scripts folder
           - script_drafts.json -> refined_scripts.json (rename for S5 convention)
        3. If character_registry.json doesn't exist in S4, create empty one in S5
        """
        import shutil
        import glob
        
        project_dir = self.projects.get_project_path(project_name)
        s4_dir = os.path.join(project_dir, "4_scripts")
        s5_dir = os.path.join(project_dir, "5_scripts")
        
        # Check S4 has data
        s4_scripts_path = os.path.join(s4_dir, "script_drafts.json")
        if not os.path.exists(s4_scripts_path):
            raise ValueError("Stage 4 无数据可拷贝")
        
        with open(s4_scripts_path, 'r', encoding='utf-8') as f:
            s4_data = json.load(f)
        if not s4_data:
            raise ValueError("Stage 4 无数据可拷贝")
        
        # Step 1: Clear S5 folder
        if os.path.exists(s5_dir):
            for file in glob.glob(os.path.join(s5_dir, "*")):
                if os.path.isfile(file):
                    os.remove(file)
        
        # Step 2: Ensure S5 folder exists
        os.makedirs(s5_dir, exist_ok=True)
        
        # Step 3: Copy all files from S4 to S5 with proper renaming
        copied_files = []
        # File name mapping: S4 filename -> S5 filename
        rename_map = {
            "script_drafts.json": "refined_scripts.json"  # S4 uses script_drafts, S5 uses refined_scripts
        }
        
        if os.path.exists(s4_dir):
            for file in glob.glob(os.path.join(s4_dir, "*")):
                if os.path.isfile(file):
                    filename = os.path.basename(file)
                    # Apply rename mapping if exists
                    dest_filename = rename_map.get(filename, filename)
                    dest = os.path.join(s5_dir, dest_filename)
                    shutil.copy2(file, dest)
                    if filename != dest_filename:
                        copied_files.append(f"{filename} -> {dest_filename}")
                    else:
                        copied_files.append(filename)
        
        # Step 4: Ensure character_registry.json exists
        s5_registry_path = os.path.join(s5_dir, "character_registry.json")
        if not os.path.exists(s5_registry_path):
            with open(s5_registry_path, 'w', encoding='utf-8') as f:
                json.dump({}, f, indent=2, ensure_ascii=False)
            copied_files.append("character_registry.json (created empty)")
        
        return {"success": True, "count": len(s4_data), "copied_files": copied_files}

    def check_needs_init(self, project_name: str) -> Dict[str, Any]:
        """Check if Stage 5 needs initialization from Stage 4."""
        project_dir = self.projects.get_project_path(project_name)
        s4_path = os.path.join(project_dir, "4_scripts/script_drafts.json")
        s5_path = os.path.join(project_dir, "5_scripts/refined_scripts.json")
        
        s4_data = []
        s5_data = []
        
        if os.path.exists(s4_path):
            with open(s4_path, 'r', encoding='utf-8') as f:
                s4_data = json.load(f)
        
        if os.path.exists(s5_path):
            with open(s5_path, 'r', encoding='utf-8') as f:
                s5_data = json.load(f)
        
        return {
            "needs_init": not os.path.exists(s5_path) or len(s5_data) == 0,
            "s4_count": len(s4_data),
            "s5_count": len(s5_data)
        }

